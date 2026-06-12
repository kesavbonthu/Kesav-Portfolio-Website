from __future__ import annotations

import json
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
NODE_BIN = Path.home() / ".cache/codex-runtimes/codex-primary-runtime/dependencies/node/bin/node"
OUTPUT_DOCX = ROOT / "Kesav Portfolio Website Content.docx"


def load_site_data() -> dict:
    node_script = r"""
const fs = require("fs");
const path = require("path");
const vm = require("vm");
const ts = require("typescript");

const root = process.argv[1];

function loadTsModule(relPath) {
  const filePath = path.join(root, relPath);
  const source = fs.readFileSync(filePath, "utf8");
  const transpiled = ts.transpileModule(source, {
    compilerOptions: {
      module: ts.ModuleKind.CommonJS,
      target: ts.ScriptTarget.ES2020,
    },
    fileName: filePath,
  }).outputText;

  const module = { exports: {} };
  const context = vm.createContext({
    module,
    exports: module.exports,
    require,
    __dirname: path.dirname(filePath),
    __filename: filePath,
    console,
    process,
  });

  new vm.Script(transpiled, { filename: filePath }).runInContext(context);
  return module.exports;
}

const appPage = fs.readFileSync(path.join(root, "app/page.tsx"), "utf8");

function extractArray(name) {
  const match = appPage.match(new RegExp(`const ${name} = \\[(.*?)\\];`, "s"));
  if (!match) return [];
  return Function(`return [${match[1]}];`)();
}

const profile = loadTsModule("data/profile.ts").profile;
const experiences = loadTsModule("data/experience.ts").experiences;
const projects = loadTsModule("data/projects.ts").projects;
const consultingProjects = loadTsModule("data/consulting.ts").consultingProjects;
const skillCategories = loadTsModule("data/skills.ts").skillCategories;

const featuredCaseStudyOrder = extractArray("featuredCaseStudyOrder");
const aiProjectSlugs = extractArray("aiProjectSlugs");
const whatIBring = extractArray("whatIBring");
const targetRoles = extractArray("targetRoles");

console.log(JSON.stringify({
  profile,
  experiences,
  projects,
  consultingProjects,
  skillCategories,
  featuredCaseStudyOrder,
  aiProjectSlugs,
  whatIBring,
  targetRoles
}));
"""

    result = subprocess.run(
        [str(NODE_BIN), "-e", node_script, str(ROOT)],
        check=True,
        capture_output=True,
        text=True,
    )
    return json.loads(result.stdout)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color in [
        ("Title", 24, RGBColor(17, 24, 39)),
        ("Heading 1", 16, RGBColor(15, 23, 42)),
        ("Heading 2", 13, RGBColor(30, 41, 59)),
        ("Heading 3", 11.5, RGBColor(51, 65, 85)),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color


def add_title_block(doc: Document, profile: dict) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Kesav Portfolio Website Content")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(17, 24, 39)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(
        "Editable export of the current portfolio website for bulk content updates"
    )
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(10.5)
    subtitle_run.font.color.rgb = RGBColor(71, 85, 105)

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(1.85)
    table.columns[1].width = Inches(4.95)
    rows = [
        ("Name", profile["name"]),
        ("Current site", "https://kesavbonthu.github.io"),
        ("Primary email", profile["email"]),
        ("LinkedIn", profile["linkedin"]),
    ]
    for row, (label, value) in zip(table.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], "E2E8F0")
    doc.add_paragraph(
        "Update any text directly in this file and send it back when you want the website refreshed"
    )


def add_section_heading(doc: Document, title: str, body: str | None = None) -> None:
    doc.add_paragraph(title, style="Heading 1")
    if body:
        doc.add_paragraph(body)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_inline_label_paragraph(doc: Document, label: str, value: str) -> None:
    paragraph = doc.add_paragraph()
    label_run = paragraph.add_run(f"{label}: ")
    label_run.bold = True
    value_run = paragraph.add_run(value)
    value_run.font.name = "Arial"


def add_project_block(doc: Document, project: dict) -> None:
    doc.add_paragraph(project["title"], style="Heading 2")
    meta_parts = []
    if project.get("company"):
        meta_parts.append(project["company"])
    if project.get("timeframe"):
        meta_parts.append(project["timeframe"])
    if project.get("categories"):
        meta_parts.append(", ".join(project["categories"]))
    if meta_parts:
        meta = doc.add_paragraph(" | ".join(meta_parts))
        meta.runs[0].italic = True

    add_inline_label_paragraph(doc, "Summary", project["summary"])
    add_inline_label_paragraph(doc, "Problem", project["problem"])
    add_inline_label_paragraph(doc, "Users", ", ".join(project["users"]))
    add_inline_label_paragraph(doc, "Role", project["role"])
    doc.add_paragraph("Approach", style="Heading 3")
    add_bullets(doc, project["approach"])
    doc.add_paragraph("Solution", style="Heading 3")
    add_bullets(doc, project["solution"])
    doc.add_paragraph("Impact", style="Heading 3")
    add_bullets(doc, project["impact"])
    add_inline_label_paragraph(doc, "Skills", ", ".join(project["skills"]))
    if project.get("lessonsLearned"):
        add_inline_label_paragraph(doc, "Lessons learned", project["lessonsLearned"])


def add_consulting_block(doc: Document, project: dict) -> None:
    doc.add_paragraph(project["project"], style="Heading 2")
    meta = doc.add_paragraph(project["client"])
    meta.runs[0].italic = True
    add_inline_label_paragraph(doc, "Context", project["context"])
    add_inline_label_paragraph(doc, "Challenge", project["challenge"])
    doc.add_paragraph("Approach", style="Heading 3")
    add_bullets(doc, project["approach"])
    doc.add_paragraph("Recommendations", style="Heading 3")
    add_bullets(doc, project["recommendations"])
    doc.add_paragraph("Impact", style="Heading 3")
    add_bullets(doc, project["impact"])
    add_inline_label_paragraph(doc, "Skills", ", ".join(project["skills"]))


def add_experience_block(doc: Document, exp: dict) -> None:
    doc.add_paragraph(f'{exp["company"]} | {exp["role"]}', style="Heading 2")
    meta = doc.add_paragraph(f'{exp["startDate"]} - {exp["endDate"]} | {exp["location"]}')
    meta.runs[0].italic = True
    add_inline_label_paragraph(doc, "Summary", exp["summary"])
    doc.add_paragraph("Highlights", style="Heading 3")
    add_bullets(doc, exp["highlights"])
    if exp["skills"]:
        add_inline_label_paragraph(doc, "Skills", ", ".join(exp["skills"]))


def build_document(data: dict) -> Document:
    doc = Document()
    style_document(doc)
    add_title_block(doc, data["profile"])

    add_section_heading(doc, "Hero Section")
    add_inline_label_paragraph(doc, "Display name", data["profile"]["name"])
    add_inline_label_paragraph(
        doc,
        "Main headline",
        "Product and Program Manager building AI, financial services, and enterprise workflow solutions",
    )
    add_inline_label_paragraph(
        doc,
        "Intro paragraph",
        "Recent MBA grad at UW Foster with 6+ years of experience across product management, technical program execution, and software engineering. I specialize in building data-driven products, managing complex cross-functional initiatives, and turning technical complexity into measurable business outcomes",
    )
    add_inline_label_paragraph(doc, "Credential badges", ", ".join(data["profile"]["credentials"]))
    add_inline_label_paragraph(doc, "Primary CTA", "Connect on LinkedIn")

    add_section_heading(
        doc,
        "About / Positioning",
        "Product and Program Manager with engineering depth, AI experience, and a track record of measurable impact",
    )
    add_paragraphs = [
        data["profile"]["summary"],
        "I combine product strategy, technical program execution, and hands-on engineering depth. My experience spans software engineering and technical product roles at Citi, AI product management at EndoMD Health, and MBA consulting work at UW Foster across enterprise process design and payments strategy",
        "I bring a blend of product thinking, program execution, and technical fluency to teams building data-driven products, leading complex cross-functional initiatives, and turning technical complexity into business outcomes",
    ]
    for paragraph in add_paragraphs:
        doc.add_paragraph(paragraph)

    doc.add_paragraph("What I Bring", style="Heading 2")
    for item in data["whatIBring"]:
        add_inline_label_paragraph(doc, item["title"], item["body"])
    add_inline_label_paragraph(
        doc,
        "Core differentiator",
        "Equally strong in product strategy and program execution, with the technical fluency to work credibly across engineering, operations, analytics, and executive stakeholders",
    )

    featured_lookup = {project["slug"]: project for project in data["projects"]}
    add_section_heading(
        doc,
        "Featured Case Studies",
        "Product, program, AI, financial services, and enterprise execution in one portfolio",
    )
    for slug in data["featuredCaseStudyOrder"]:
        project = featured_lookup.get(slug)
        if project:
            add_project_block(doc, project)

    add_section_heading(
        doc,
        "AI Projects",
        "Personal AI and agentic workflow work beyond core operating roles",
    )
    for slug in data["aiProjectSlugs"]:
        project = featured_lookup.get(slug)
        if project:
            add_project_block(doc, project)

    add_section_heading(
        doc,
        "Consulting Work",
        "MBA strategy work across enterprise workflow design and payments growth",
    )
    for project in data["consultingProjects"]:
        add_consulting_block(doc, project)

    experiences = data["experiences"]
    add_section_heading(doc, "Career Experience")
    doc.add_paragraph("Industry Roles", style="Heading 2")
    for exp in [item for item in experiences if item["type"] == "full-time"]:
        add_experience_block(doc, exp)
    doc.add_paragraph("MBA Consulting Projects", style="Heading 2")
    for exp in [item for item in experiences if item["type"] == "consulting"]:
        add_experience_block(doc, exp)
    doc.add_paragraph("Education", style="Heading 2")
    for exp in [item for item in experiences if item["type"] == "education"]:
        add_experience_block(doc, exp)

    add_section_heading(doc, "Skills")
    for category in data["skillCategories"]:
        if category["category"] not in {
            "Product Management",
            "Technical Program Management",
            "AI / LLM / Agentic AI",
            "Engineering",
        }:
            continue
        doc.add_paragraph(category["category"], style="Heading 2")
        add_bullets(doc, category["skills"])

    add_section_heading(
        doc,
        "Target Roles",
        "Interested in product, technical product, technical program and program manager roles",
    )
    for role in data["targetRoles"]:
        add_inline_label_paragraph(doc, role["title"], role["body"])

    add_section_heading(
        doc,
        "Contact",
        "Open to conversations around product, program, AI, and enterprise platform work",
    )
    add_inline_label_paragraph(doc, "Email", data["profile"]["email"])
    add_inline_label_paragraph(doc, "LinkedIn", data["profile"]["linkedin"])
    add_inline_label_paragraph(doc, "GitHub", data["profile"]["github"])
    add_inline_label_paragraph(doc, "Location", data["profile"]["location"])
    add_inline_label_paragraph(doc, "Availability", "Open to on-site, hybrid, and remote roles")
    add_inline_label_paragraph(doc, "Relocation", "Open to relocation within US")

    return doc


def main() -> None:
    data = load_site_data()
    document = build_document(data)
    document.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    try:
        main()
    except subprocess.CalledProcessError as exc:
        sys.stderr.write(exc.stderr or str(exc))
        raise
